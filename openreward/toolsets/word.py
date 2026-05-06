"""使用 python-docx 的 Word 文档操作工具集。"""

from __future__ import annotations

import json
from typing import Literal, Optional, Any
from pydantic import BaseModel, Field

from openreward.environments.environment import tool
from openreward.environments.toolset import Toolset
from openreward.environments.types import TextBlock, ToolOutput


# ===== Pydantic 参数模型 =====

class CreateDocumentParams(BaseModel):
    file_path: str = Field(..., description="新文档的创建路径")
    title: str | None = Field(None, description="文档标题元数据")
    author: str | None = Field(None, description="文档作者元数据")


class GetOverviewParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")


class ReadDocumentParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    include_tables: bool = Field(True, description="在输出中包含表格内容")
    max_paragraphs: int | None = Field(None, description="限制返回的段落数量")


class ReadImageParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    image_index: int | None = Field(None, description="要读取的特定图像索引（如果为 None，则返回所有图像）")
    include_data: bool = Field(True, description="在响应中包含 base64 编码的图像数据")


class AddContentParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    text: str = Field(..., description="要添加的文本内容")
    content_type: Literal["paragraph", "heading"] = Field("paragraph", description="内容类型")
    heading_level: int | None = Field(None, description="如果 content_type 为 heading，则为标题级别（1-9）")
    style: str | None = Field(None, description="段落样式名称")
    position: Literal["end", "start"] = Field("end", description="插入内容的位置")


class AddImageParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    image_path: str = Field(..., description="要插入的图像文件路径")
    width_inches: float | None = Field(None, description="图像宽度（英寸）")
    height_inches: float | None = Field(None, description="图像高度（英寸）")
    position: Literal["end", "start"] = Field("end", description="插入图像的位置")


class EditContentParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    paragraph_index: int = Field(..., description="从零开始的段落索引")
    new_text: str = Field(..., description="新文本内容")
    append: bool = Field(False, description="追加到现有文本而不是替换")


class DeleteContentParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    paragraph_index: int = Field(..., description="要删除的从零开始的段落索引")


class ModifyImageParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    image_index: int = Field(..., description="从零开始的图像索引")
    width_inches: float | None = Field(None, description="新宽度（英寸）")
    height_inches: float | None = Field(None, description="新高度（英寸）")


class ApplyFormattingParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    paragraph_index: int = Field(..., description="从零开始的段落索引")
    bold: bool | None = Field(None, description="应用粗体格式")
    italic: bool | None = Field(None, description="应用斜体格式")
    underline: bool | None = Field(None, description="应用下划线格式")
    font_name: str | None = Field(None, description="字体族名称（例如，'Arial'）")
    font_size: int | None = Field(None, description="字体大小（磅）")
    color: str | None = Field(None, description="十六进制颜色代码（例如，'FF0000' 表示红色）")


class DeleteDocumentParams(BaseModel):
    file_path: str = Field(..., description="要删除的 Word 文档路径")


class SearchDocumentParams(BaseModel):
    file_path: str = Field(..., description="Word 文档路径")
    search_text: str = Field(..., description="要搜索的文本")
    case_sensitive: bool = Field(False, description="区分大小写的搜索")
    search_in_tables: bool = Field(True, description="在搜索中包含表格内容")
    max_results: int | None = Field(None, description="限制结果数量")


# ===== Word 工具集类 =====

class WordToolset(Toolset):
    """通过 python-docx 在沙箱中执行，提供 Word 文档操作工具的工具集"""

    @tool
    async def word_create_document(self, params: CreateDocumentParams) -> ToolOutput:
        """创建带有可选元数据的新 Word 文档"""
        title_escaped = params.title.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n') if params.title else ""
        author_escaped = params.author.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n') if params.author else ""

        script = f'''
import json
from docx import Document

try:
    doc = Document()

    # 如果提供了元数据则设置
    if "{title_escaped}":
        doc.core_properties.title = "{title_escaped}"
    if "{author_escaped}":
        doc.core_properties.author = "{author_escaped}"

    doc.save("{params.file_path}")

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "title": "{title_escaped}" if "{title_escaped}" else None,
        "author": "{author_escaped}" if "{author_escaped}" else None,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Document created successfully at {params.file_path}")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_delete_document(self, params: DeleteDocumentParams) -> ToolOutput:
        """从沙箱中删除 Word 文档文件"""
        output, exit_code = await self.sandbox.run(f"rm {params.file_path}", max_bytes=1_000_000)

        if exit_code == 0:
            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Document deleted successfully: {params.file_path}")],
                metadata={"file_path": params.file_path, "deleted": True},
                reward=0.0,
                finished=False,
            )
        else:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Failed to delete: {output}")],
                metadata={"file_path": params.file_path, "deleted": False, "error": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_get_document_overview(self, params: GetOverviewParams) -> ToolOutput:
        """检索文档结构和元数据，包括段落数、表格数和图像数"""
        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    # 统计段落数
    paragraph_count = len(doc.paragraphs)

    # 统计表格数
    table_count = len(doc.tables)

    # 统计图像数（内联形状）
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    # 获取元数据
    props = doc.core_properties
    metadata = {{
        "title": props.title if props.title else "",
        "author": props.author if props.author else "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
    }}

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "image_count": image_count,
        "metadata": metadata,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            # 构建摘要
            summary = f"Document: {params.file_path}\n"
            summary += f"Paragraphs: {result['paragraph_count']}\n"
            summary += f"Tables: {result['table_count']}\n"
            summary += f"Images: {result['image_count']}\n\n"
            summary += "Metadata:\n"
            for key, value in result['metadata'].items():
                if value:
                    summary += f"  {key.capitalize()}: {value}\n"

            return ToolOutput(
                blocks=[TextBlock(text=summary)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_read_document_content(self, params: ReadDocumentParams) -> ToolOutput:
        """从文档中提取所有带有段落结构的文本内容"""
        max_para = params.max_paragraphs if params.max_paragraphs else 999999

        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        if idx >= {max_para}:
            break

        paragraphs.append({{
            "index": idx,
            "text": para.text,
            "style": para.style.name if para.style else "Normal",
        }})

    # 可选地包含表格内容
    tables = []
    if {params.include_tables}:
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({{
                "index": table_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data,
            }})

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "table_count": len(tables),
        "tables": tables,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            # 构建摘要
            summary = f"Document content from {params.file_path}\n"
            summary += f"Paragraphs: {result['paragraph_count']}\n"
            if result['table_count'] > 0:
                summary += f"Tables: {result['table_count']}\n"
            summary += "\n"

            # 显示前几个段落
            for para in result['paragraphs']:
                text_preview = para['text'] if para['text'] else "(empty)"
                summary += f"{text_preview}\n\n"

            return ToolOutput(
                blocks=[TextBlock(text=summary)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_read_image(self, params: ReadImageParams) -> ToolOutput:
        """从文档中提取图像数据和元数据"""
        script = f'''
import json
import base64
from docx import Document

try:
    doc = Document("{params.file_path}")

    images = []

    # 遍历所有部分以查找图像
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_data = image_part.blob  # 二进制数据

            # 从 content_type 提取格式（例如，'image/png' -> 'png'）
            image_format = image_part.content_type.split('/')[-1]

            image_info = {{
                "index": len(images),
                "format": image_format,
                "size_bytes": len(image_data),
            }}

            # 如果请求则包含 base64 数据
            if {params.include_data}:
                image_base64 = base64.b64encode(image_data).decode('utf-8')
                image_info["data"] = image_base64

            images.append(image_info)

    # 如果指定了 image_index 则过滤
    if {params.image_index} is not None:
        if {params.image_index} < 0 or {params.image_index} >= len(images):
            print(json.dumps({{
                "success": False,
                "error": f"Image index {params.image_index} out of range (document has {{len(images)}} images)"
            }}))
        else:
            images = [images[{params.image_index}]]

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "image_count": len(images),
        "images": images,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            # 构建摘要
            summary = f"Found {result['image_count']} image(s) in {params.file_path}\n\n"
            for img in result['images']:
                summary += f"Image {img['index']}: {img['format'].upper()} ({img['size_bytes']} bytes)\n"
                if params.include_data and 'data' in img:
                    summary += f"  Base64 data: {len(img['data'])} characters\n"

            return ToolOutput(
                blocks=[TextBlock(text=summary)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_add_content_text(self, params: AddContentParams) -> ToolOutput:
        """向文档添加段落或标题"""
        text_escaped = params.text.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')
        style_escaped = params.style.replace('\\', '\\\\').replace('"', '\\"') if params.style else ""

        # 构建内容添加逻辑
        if params.content_type == "heading":
            if not params.heading_level or params.heading_level < 1 or params.heading_level > 9:
                return ToolOutput(
                    blocks=[TextBlock(text="❌ Error: heading_level must be between 1 and 9 for heading content")],
                    metadata={"error": "Invalid heading_level"},
                    reward=0.0,
                    finished=False,
                )
            add_code = f'new_para = doc.add_heading("{text_escaped}", level={params.heading_level})'
        else:
            add_code = f'new_para = doc.add_paragraph("{text_escaped}")'
            if params.style:
                add_code += f'\n    new_para.style = "{style_escaped}"'

        # 处理位置（开头与末尾）
        if params.position == "start":
            insert_code = f'''
    # 通过在第一个段落之前插入来添加到开头
    if len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]._element
        p.getparent().insert(0, new_para._element)
    else:
        {add_code}
'''
        else:
            insert_code = f'    {add_code}'

        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

{insert_code}

    doc.save("{params.file_path}")

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "content_type": "{params.content_type}",
        "position": "{params.position}",
        "paragraph_count": len(doc.paragraphs),
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            content_label = f"Heading {params.heading_level}" if params.content_type == "heading" else "Paragraph"
            return ToolOutput(
                blocks=[TextBlock(text=f"✅ {content_label} added at {params.position} of document")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_add_image(self, params: AddImageParams) -> ToolOutput:
        """向文档插入图像"""
        width_param = f", width=Inches({params.width_inches})" if params.width_inches else ""
        height_param = f", height=Inches({params.height_inches})" if params.height_inches else ""

        # 处理位置
        if params.position == "start":
            insert_code = f'''
    # 在开头的新段落中添加图像
    paragraph = doc.paragraphs[0] if len(doc.paragraphs) > 0 else doc.add_paragraph()
    run = paragraph.insert_paragraph_before().add_run()
    picture = run.add_picture("{params.image_path}"{width_param}{height_param})
'''
        else:
            insert_code = f'''
    # 在末尾添加图像
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    picture = run.add_picture("{params.image_path}"{width_param}{height_param})
'''

        script = f'''
import json
from docx import Document
from docx.shared import Inches

try:
    doc = Document("{params.file_path}")

{insert_code}

    # 获取最终尺寸
    width_inches = picture.width / 914400
    height_inches = picture.height / 914400

    doc.save("{params.file_path}")

    result = {{
        "success": True,
        "file_path": "{params.file_path}",
        "image_path": "{params.image_path}",
        "position": "{params.position}",
        "width": width_inches,
        "height": height_inches,
    }}
    print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Image added at {params.position} of document ({result['width']:.2f}\" x {result['height']:.2f}\")")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_edit_content_text(self, params: EditContentParams) -> ToolOutput:
        """修改现有段落文本"""
        text_escaped = params.new_text.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')

        if params.append:
            update_code = f'para.add_run("\\n{text_escaped}")'
        else:
            update_code = f'para.text = "{text_escaped}"'

        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    if {params.paragraph_index} < 0 or {params.paragraph_index} >= len(doc.paragraphs):
        print(json.dumps({{
            "success": False,
            "error": f"Paragraph index {params.paragraph_index} out of range (document has {{len(doc.paragraphs)}} paragraphs)"
        }}))
    else:
        para = doc.paragraphs[{params.paragraph_index}]
        old_text = para.text

        {update_code}

        doc.save("{params.file_path}")

        result = {{
            "success": True,
            "file_path": "{params.file_path}",
            "paragraph_index": {params.paragraph_index},
            "old_text": old_text[:100],
            "new_text": para.text[:100],
            "append": {params.append},
        }}
        print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            action = "appended to" if params.append else "updated"
            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Paragraph {params.paragraph_index} {action} successfully\nPreview: {result['new_text']}...")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_delete_content_text(self, params: DeleteContentParams) -> ToolOutput:
        """按索引删除段落"""
        script = f'''
import json
from docx import Document

try:
    doc = Document("{params.file_path}")

    if {params.paragraph_index} < 0 or {params.paragraph_index} >= len(doc.paragraphs):
        print(json.dumps({{
            "success": False,
            "error": f"Paragraph index {params.paragraph_index} out of range (document has {{len(doc.paragraphs)}} paragraphs)"
        }}))
    else:
        para = doc.paragraphs[{params.paragraph_index}]
        deleted_text = para.text

        # 删除段落元素
        p = para._element
        p.getparent().remove(p)

        doc.save("{params.file_path}")

        result = {{
            "success": True,
            "file_path": "{params.file_path}",
            "paragraph_index": {params.paragraph_index},
            "deleted_text": deleted_text[:100],
            "remaining_paragraphs": len(doc.paragraphs),
        }}
        print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Paragraph {params.paragraph_index} deleted successfully\nDeleted text: {result['deleted_text']}...")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_modify_image(self, params: ModifyImageParams) -> ToolOutput:
        """更新图像属性（大小）"""
        width_code = f"picture.width = Inches({params.width_inches})" if params.width_inches else ""
        height_code = f"picture.height = Inches({params.height_inches})" if params.height_inches else ""

        script = f'''
import json
from docx import Document
from docx.shared import Inches

try:
    doc = Document("{params.file_path}")

    # 如果可用则尝试使用 inline_shapes
    if hasattr(doc, 'inline_shapes'):
        if {params.image_index} < 0 or {params.image_index} >= len(doc.inline_shapes):
            print(json.dumps({{
                "success": False,
                "error": f"Image index {params.image_index} out of range (document has {{len(doc.inline_shapes)}} images)"
            }}))
        else:
            picture = doc.inline_shapes[{params.image_index}]

            {width_code}
            {height_code}

            new_width = picture.width / 914400
            new_height = picture.height / 914400

            doc.save("{params.file_path}")

            result = {{
                "success": True,
                "file_path": "{params.file_path}",
                "image_index": {params.image_index},
                "new_width": new_width,
                "new_height": new_height,
            }}
            print(json.dumps(result))
    else:
        print(json.dumps({{
            "success": False,
            "error": "Image modification requires python-docx with inline_shapes support"
        }}))

except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Image {params.image_index} modified successfully ({result['new_width']:.2f}\" x {result['new_height']:.2f}\")")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_apply_formatting(self, params: ApplyFormattingParams) -> ToolOutput:
        """将文本格式应用于段落 run"""
        # 构建格式代码
        format_lines = []
        if params.bold is not None:
            format_lines.append(f"run.font.bold = {str(params.bold)}")
        if params.italic is not None:
            format_lines.append(f"run.font.italic = {str(params.italic)}")
        if params.underline is not None:
            format_lines.append(f"run.font.underline = {str(params.underline)}")
        if params.font_name:
            font_escaped = params.font_name.replace('\\', '\\\\').replace('"', '\\"')
            format_lines.append(f'run.font.name = "{font_escaped}"')
        if params.font_size:
            format_lines.append(f"run.font.size = Pt({params.font_size})")
        if params.color:
            # 将十六进制颜色（例如，"FF0000"）转换为 RGB
            format_lines.append(f"run.font.color.rgb = RGBColor(int('{params.color}'[0:2], 16), int('{params.color}'[2:4], 16), int('{params.color}'[4:6], 16))")

        format_code = "\n                ".join(format_lines) if format_lines else "pass"

        script = f'''
import json
from docx import Document
from docx.shared import Pt, RGBColor

try:
    doc = Document("{params.file_path}")

    if {params.paragraph_index} < 0 or {params.paragraph_index} >= len(doc.paragraphs):
        print(json.dumps({{
            "success": False,
            "error": f"Paragraph index {params.paragraph_index} out of range (document has {{len(doc.paragraphs)}} paragraphs)"
        }}))
    else:
        para = doc.paragraphs[{params.paragraph_index}]

        # 将格式应用于段落中的所有 run
        for run in para.runs:
            if run.text.strip():
                {format_code}

        doc.save("{params.file_path}")

        result = {{
            "success": True,
            "file_path": "{params.file_path}",
            "paragraph_index": {params.paragraph_index},
            "formatting_applied": {{
                "bold": {params.bold},
                "italic": {params.italic},
                "underline": {params.underline},
                "font_name": "{params.font_name}" if "{params.font_name}" else None,
                "font_size": {params.font_size},
                "color": "{params.color}" if "{params.color}" else None,
            }},
        }}
        print(json.dumps(result))
except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            return ToolOutput(
                blocks=[TextBlock(text=f"✅ Formatting applied to paragraph {params.paragraph_index} successfully")],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )

    @tool
    async def word_search_document(self, params: SearchDocumentParams) -> ToolOutput:
        """在文档段落和表格中搜索文本"""
        search_text_escaped = params.search_text.replace('\\', '\\\\').replace('"', '\\"')
        max_results_repr = repr(params.max_results)

        script = f'''
import json
from docx import Document
import re

try:
    doc = Document("{params.file_path}")

    # 准备搜索模式
    flags = 0 if {params.case_sensitive} else re.IGNORECASE
    pattern = re.compile(re.escape("{search_text_escaped}"), flags=flags)

    results = []
    max_results = {max_results_repr}

    # 在段落中搜索
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text:
            continue

        for match in pattern.finditer(text):
            start = max(0, match.start() - 50)
            end = min(len(text), match.end() + 50)
            context = text[start:end]

            results.append({{
                "location_type": "paragraph",
                "paragraph_index": para_idx,
                "start_pos": match.start(),
                "end_pos": match.end(),
                "context": context,
                "style": para.style.name if para.style else "Normal"
            }})

            if max_results and len(results) >= max_results:
                break

        if max_results and len(results) >= max_results:
            break

    # 如果启用了表格搜索且未达到 max_results
    if {params.search_in_tables} and (not max_results or len(results) < max_results):
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text
                    if not text:
                        continue

                    for match in pattern.finditer(text):
                        start = max(0, match.start() - 50)
                        end = min(len(text), match.end() + 50)
                        context = text[start:end]

                        results.append({{
                            "location_type": "table",
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "col_index": col_idx,
                            "start_pos": match.start(),
                            "end_pos": match.end(),
                            "context": context
                        }})

                        if max_results and len(results) >= max_results:
                            break

                    if max_results and len(results) >= max_results:
                        break

                if max_results and len(results) >= max_results:
                    break

            if max_results and len(results) >= max_results:
                break

    result = {{
        "success": True,
        "search_text": "{search_text_escaped}",
        "total_results": len(results),
        "results": results
    }}
    print(json.dumps(result))

except Exception as e:
    print(json.dumps({{"success": False, "error": str(e)}}))
'''

        cmd = f"python3 << 'PYSCRIPT_EOF'\n{script}\nPYSCRIPT_EOF"
        output, exit_code = await self.sandbox.run(cmd, max_bytes=1_000_000)
        try:
            result = json.loads(output)

            if not result.get("success"):
                return ToolOutput(
                    blocks=[TextBlock(text=f"❌ Error: {result.get('error')}")],
                    metadata={"error": result.get("error")},
                    reward=0.0,
                    finished=False,
                )

            count = result["total_results"]
            display_text = f"✅ Found {count} occurrence(s) of '{params.search_text}'"

            if count > 0:
                display_text += "\n\nMatches:"
                for i, match in enumerate(result["results"][:5], 1):
                    if match["location_type"] == "paragraph":
                        display_text += f"\n{i}. Paragraph {match['paragraph_index']} ({match['style']}): ...{match['context']}..."
                    else:
                        display_text += f"\n{i}. Table {match['table_index']} [Row {match['row_index']}, Col {match['col_index']}]: ...{match['context']}..."

                if count > 5:
                    display_text += f"\n\n... and {count - 5} more"

            return ToolOutput(
                blocks=[TextBlock(text=display_text)],
                metadata=result,
                reward=0.0,
                finished=False,
            )
        except json.JSONDecodeError:
            return ToolOutput(
                blocks=[TextBlock(text=f"❌ Error parsing output: {output}")],
                metadata={"error": "JSON decode failed", "output": output},
                reward=0.0,
                finished=False,
            )
